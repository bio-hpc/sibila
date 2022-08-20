#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""MergeResults.py:
    Gather all the information from the json files obtained in the tests and create a table in a csv file
"""
__author__ = "Jorge de la Peña García"
__version__ = "1.0"
__maintainer__ = "Jorge"
__email__ = "jpena@ucam.edu"
__status__ = "Production"

import json
from glob import glob
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx import Document
from docx.shared import Inches, Pt
from os.path import dirname, basename
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from Tools.Word2Pdf import convert_to

GRAPHS = ['roc_proba', 'confusion_matrix', 'correlation', 'PermutationImportance_hist', 'Lime', 'IntegratedGradients', 'Shapley', 'Dice', 'times']
LST_PREFIX = ["ALE", "PDP"]


class JoinGraphs:
    def __init__(self, dir_name):
        if dir_name[:-1] != "/":
            dir_name += "/"
        self.dir_name = dir_name
        self.file_out = dir_name + basename(dirname(dir_name)) + '.docx'
        self.file_pdf = dir_name + basename(dirname(dir_name)) + '.pdf'

    def get_number_rows(self, lst_files):
        number_rows = 0
        for name in GRAPHS:
            number_rows += len(self._search_files(lst_files, name))
        return number_rows

    def paragraph_format_run(self, cell):
        paragraph = cell.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = paragraph.add_run()
        return run

    def create_document(self):
        self.doc = Document()
        self.doc.add_heading(dirname(self.dir_name), 0)

    def create_doc(self, method, save_doc=True):
        lst_files = sorted(glob(self.dir_name + "/" + method + "*.png"))
        table = self.doc.add_table(rows=self.get_number_rows(lst_files) + 1, cols=1)
        table.allow_autofit = False
        self.create_first_row(method, table)
        #cell = table.cell(0, 0)
        #cell.text = method

        cnt_row = 1
        for name in GRAPHS:
            for file in self._search_files(lst_files, name):
                cell_r = self.paragraph_format_run(table.cell(cnt_row, 0))
                cell_r.add_picture(file, width=Inches(5.0))
                cnt_row += 1

        if save_doc:
            self.save()

    def _search_files(self, lst_files, name):
        return sorted([i for i in lst_files if name in i])

    def save(self, file_name=None):
        self.doc.add_page_break()
        if not file_name:
            file_name = self.file_out
        self.doc.save(file_name)
        try:
            convert_to(self.dir_name, file_name)
        except:
            pass

    def create_global(self):
        self.create_document()
        expr = self.dir_name + "[^A-Z]*_*_*.png"
        for m in self.get_methods(expr):
            self.create_doc(m, False)

        self.doc.add_page_break()
        results = self.collect_metrics()
        self.create_table_metrics(results)
        
        self.save()

    def __find_metrics(self, content):
        data = json.loads(content)
        analysis = data['Analysis']

        if 'Auc' not in analysis.keys():
            return {
                'Model': data['Config']['Model_params']['model'],
                'Pearson': [float(x) for x in analysis['Pearson Correlation Coefficient'].split('/')],
                'R2': analysis['Coefficient of Determination'],
                'MAE': analysis['Mean Absolute Error'],
                'MSE': analysis['Mean Squared Error']
            }
        else:
            return {
                'Model': data['Config']['Model_params']['model'],
                'Accuracy': analysis['Accuracy'],
                'Precision': analysis['Precision'],
                'F1': analysis['F1'],
                'Recall': analysis['Recall'],
                'Specificity': analysis['Specificity'],
                'AUC': analysis['Auc']
            }

    def get_methods(self, expr):
        lst_methods = []
        lst = glob(expr)
        for i in lst:
            m = basename(i).split("_")[0]
            if m not in lst_methods:
                lst_methods.append(m)
        return lst_methods

    def collect_metrics(self):
        lst_files = sorted(glob(self.dir_name + "/*_data.json"))
        results = {}
        for foo in lst_files:
            with open(foo, 'r') as f:
                metrics = self.__find_metrics(f.read())
                results[metrics['Model']] = metrics
                del metrics['Model']
        return results        

    def join_all_graphs(self):
        for pre in LST_PREFIX:
            self.create_document()
            expr = self.dir_name + pre + "/[^A-Z]*_*_*.png"
            for m in self.get_methods(expr):
                f1 = self.dir_name + "/" + pre + "/" + m + "*.png"
                self.create_table(m, f1)
            out_doc = self.dir_name + basename(dirname(self.dir_name)) + '_' + pre + '.docx'
            self.save(out_doc)

    def create_first_row(self, method, table):
        cell = table.cell(0, 0)
        cell.text = method
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="CCCCCC"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading_elm_1)

    def create_table(self, method, expr):
        lst_files = sorted(glob(expr))
        table = self.doc.add_table(rows=len(lst_files) + 1, cols=1)
        table.allow_autofit = False
        self.create_first_row(method, table)
        cnt_row = 1
        for file in lst_files:
            cell_r = self.paragraph_format_run(table.cell(cnt_row, 0))
            cell_r.add_picture(file, width=Inches(5.0))
            cnt_row += 1

    def create_table_metrics(self, results):
        font = self.doc.styles['Normal'].font
        font.size = Pt(8)

        n_rows = len(list(results.values())[0].keys()) + 1
        n_cols = len(results.keys()) + 1
        table = self.doc.add_table(rows=n_rows, cols=n_cols, style='Table Grid')
        table.allow_autofit = False

        hdr_cells = table.rows[0].cells
        for i, h in enumerate([' '] + list(results.keys())):
            hdr_cells[i].text = h

        for i, k in enumerate(results.keys()):
            data = results[k]
            for j, m in enumerate(data.keys()):
                table.cell(j+1, 0).text = m

                cell = table.cell(j+1, i+1)
                cell.text = str(data[m])
                p = cell.paragraphs[0]
                p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
